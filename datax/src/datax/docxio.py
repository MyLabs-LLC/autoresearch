"""Writing .docx files.

Reading is in :mod:`datax.extract` and deliberately has no third-party dependency,
because it runs over arbitrary downloaded files. Writing needs ``python-docx``, which
is only required by the ``nemotron`` source.

The contract that matters here is **round-trip fidelity**: text written by
:func:`write_text_docx` must come back byte-identical from
:func:`datax.extract.extract_docx_text`, otherwise the character offsets of gold PII
spans no longer point at the right substrings. Every layout mode below is
paragraph-based for that reason -- runs and paragraph styles change how a document
*looks* without changing the text it yields.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

# XML 1.0 forbids most C0 control characters. python-docx would raise (or emit a file
# Word refuses to open), so they are stripped before writing -- and the caller is told,
# because a stripped character shifts every offset after it.
_ILLEGAL_XML = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")

# "Label: value" at the start of a line -- rendered with a bold label in rich layout.
_LABEL_LINE = re.compile(r"^([A-Z][^:\n]{0,60}):(\s.*)?$")


class DocxWriteError(RuntimeError):
    pass


@dataclass
class WriteResult:
    path: Path
    text: str
    """The text actually written, after control-character sanitisation. Offsets should
    be computed against this, never against the caller's original string."""
    sanitized: bool
    paragraph_count: int


def _require_docx():
    try:
        import docx  # noqa: F401
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise DocxWriteError(
            "writing .docx requires python-docx; install the datax extra:\n"
            "  uv pip install -e 'datax[docx]'   (or: pip install python-docx)"
        ) from exc
    return docx


def sanitize(text: str) -> tuple[str, bool]:
    """Strip characters that cannot be represented in a .docx. Returns the cleaned
    text and whether anything was removed."""
    cleaned = _ILLEGAL_XML.sub("", text.replace("\r\n", "\n").replace("\r", "\n"))
    return cleaned, cleaned != text


def _looks_like_heading(line: str) -> bool:
    stripped = line.strip()
    if not (0 < len(stripped) <= 70):
        return False
    if stripped.endswith((".", ",", ";", ":")):
        return False
    # A heading is short, has few words, and is not a "Label: value" pair.
    return len(stripped.split()) <= 9 and not _LABEL_LINE.match(stripped)


def write_text_docx(
    path: str | Path,
    text: str,
    *,
    layout: str = "rich",
    core_properties: dict[str, str] | None = None,
) -> WriteResult:
    """Write ``text`` to a .docx, one paragraph per line.

    ``layout`` is presentation only:

    ``plain``
        Every line is a body paragraph.
    ``rich``
        Short heading-like lines get a heading style and ``Label: value`` lines get a
        bold label run. The extracted text is identical either way.

    An empty line becomes an empty paragraph, so the blank-line structure survives.
    """
    docx = _require_docx()
    if layout not in {"plain", "rich"}:
        raise DocxWriteError(f"unknown layout {layout!r}")

    clean, sanitized = sanitize(text)
    document = docx.Document()

    if core_properties:
        props = document.core_properties
        for key, value in core_properties.items():
            if hasattr(props, key):
                setattr(props, key, value)

    lines = clean.split("\n")
    for index, line in enumerate(lines):
        if layout == "rich" and index == 0 and line.strip():
            document.add_paragraph(line, style="Title")
            continue
        if layout == "rich" and _looks_like_heading(line):
            document.add_paragraph(line, style="Heading 2")
            continue
        paragraph = document.add_paragraph()
        match = _LABEL_LINE.match(line) if layout == "rich" else None
        if match and match.group(2):
            run = paragraph.add_run(f"{match.group(1)}:")
            run.bold = True
            paragraph.add_run(match.group(2))
        elif line:
            paragraph.add_run(line)

    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out))
    return WriteResult(path=out, text=clean, sanitized=sanitized, paragraph_count=len(lines))
