"""Round-trip fidelity between the .docx writer and reader.

This is the property the whole gold-label path depends on: if text written to a .docx
does not come back identical, every span offset in the gold manifest is silently
wrong. These tests are the reason to trust the offsets rather than hope.
"""

from __future__ import annotations

import pytest

from datax.docxio import sanitize, write_text_docx
from datax.extract import DocxReadError, extract_docx, extract_docx_text, looks_like_docx

pytest.importorskip("docx", reason="writing .docx requires python-docx")


SAMPLES = [
    "Simple one-line document.",
    "Line one\nLine two\nLine three",
    "Patient Name: Jane Doe\nMRN: 88213-A\n\nDiagnosis: hypertension.",
    "Tabbed\tvalues\there",
    "Blank lines\n\n\nbetween paragraphs",
    "Unicode: café, naïve, 東京, emoji \N{ROCKET}",
    "Trailing whitespace   \nand a line with only spaces\n   \nend",
    "Punctuation: quotes “Smart” and dashes — plus ampersand & angle < >",
]


@pytest.mark.parametrize("layout", ["plain", "rich"])
@pytest.mark.parametrize("text", SAMPLES)
def test_text_survives_round_trip(tmp_path, text, layout):
    path = tmp_path / "doc.docx"
    written = write_text_docx(path, text, layout=layout)
    assert extract_docx_text(path) == written.text


@pytest.mark.parametrize("layout", ["plain", "rich"])
def test_span_offsets_survive_round_trip(tmp_path, layout):
    text = "Patient: Jane Doe\nMRN: 88213-A\nSSN: 123-45-6789\nSeen on 2024-03-01."
    marks = [("Jane", "first_name"), ("88213-A", "medical_record_number"), ("123-45-6789", "ssn")]
    spans = [(text.index(s), text.index(s) + len(s), label) for s, label in marks]

    path = tmp_path / "doc.docx"
    write_text_docx(path, text, layout=layout)
    extracted = extract_docx_text(path)

    for start, end, _ in spans:
        assert extracted[start:end] == text[start:end]


def test_rich_layout_does_not_change_text(tmp_path):
    text = "Discharge Summary\nPatient Name: Jane Doe\nThis is the body of the note."
    plain = write_text_docx(tmp_path / "a.docx", text, layout="plain")
    rich = write_text_docx(tmp_path / "b.docx", text, layout="rich")
    assert plain.text == rich.text
    assert extract_docx_text(tmp_path / "a.docx") == extract_docx_text(tmp_path / "b.docx")


def test_control_characters_are_stripped_and_reported(tmp_path):
    text = "before\x07after"
    written = write_text_docx(tmp_path / "doc.docx", text)
    assert written.sanitized
    assert "\x07" not in written.text
    assert extract_docx_text(tmp_path / "doc.docx") == written.text


def test_sanitize_normalises_line_endings():
    cleaned, changed = sanitize("a\r\nb\rc")
    assert cleaned == "a\nb\nc"
    assert changed


def test_tables_are_extracted_as_tab_separated_rows(tmp_path):
    import docx

    document = docx.Document()
    document.add_paragraph("Header paragraph")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Jane Doe"
    table.cell(1, 0).text = "MRN"
    table.cell(1, 1).text = "88213-A"
    path = tmp_path / "table.docx"
    document.save(str(path))

    extracted = extract_docx(path)
    assert extracted.table_count == 1
    assert "Name\tJane Doe" in extracted.text
    assert "MRN\t88213-A" in extracted.text


def test_metadata_and_hash_are_reported(tmp_path):
    path = tmp_path / "doc.docx"
    write_text_docx(path, "Hello", core_properties={"title": "Test Doc", "author": "datax"})
    extracted = extract_docx(path)
    assert len(extracted.sha256) == 64
    assert extracted.size_bytes > 0
    assert extracted.word_count == 1
    assert extracted.core_properties.get("title") == "Test Doc"


def test_non_docx_input_is_rejected(tmp_path):
    html = tmp_path / "fake.docx"
    html.write_bytes(b"<html><body>404 Not Found</body></html>")
    assert not looks_like_docx(html)
    with pytest.raises(DocxReadError):
        extract_docx(html)


def test_zip_without_document_part_is_rejected(tmp_path):
    import zipfile

    path = tmp_path / "notword.docx"
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("hello.txt", "not a word document")
    assert not looks_like_docx(path)
    with pytest.raises(DocxReadError, match="not a Word document"):
        extract_docx(path)


def test_doctype_is_refused(tmp_path):
    """A DOCTYPE is the entry point for entity-expansion attacks; refuse outright."""
    import zipfile

    path = tmp_path / "evil.docx"
    payload = b'<?xml version="1.0"?><!DOCTYPE r [<!ENTITY a "x">]><w:document/>'
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("word/document.xml", payload)
    with pytest.raises(DocxReadError, match="DOCTYPE"):
        extract_docx(path)
